import io
import re

import fitz
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm
from docxcompose.composer import Composer

ID_NAME_RE = re.compile(r"^(\d{5})\s+(.+)$")
ID_ONLY_RE = re.compile(r"^(\d{5})$")
PLACEHOLDER_RE = re.compile(r"\{\{(이름|학번|반|번호)\}\}")
CROP_PAD = 4
DEFAULT_QR_WIDTH_CM = 4.5


def _parse_cell_text(text_block):
    """반환: (학번, 이름, 문제사유). 이름을 못찾으면 이름=None과 사유 문자열을 함께 반환한다."""
    lines = text_block["lines"]
    full_text = " ".join("".join(s["text"] for s in l["spans"]) for l in lines)
    full_text = full_text.replace("\xa0", " ")
    full_text = re.sub(r"\s+", " ", full_text).strip()

    m = ID_NAME_RE.match(full_text)
    if m:
        return m.group(1), m.group(2), None

    m = ID_ONLY_RE.match(full_text)
    if m:
        return m.group(1), None, "이름이 비어 있습니다"

    return None, None, f"학번/이름 형식을 인식하지 못했습니다: {full_text!r}"


def parse_source_pdf(pdf_bytes):
    """QR 그리드 PDF를 파싱한다.

    반환: (students, issues)
    - students: 정상 인식된 학생 dict 목록 (학번/이름/반/번호/크롭된 QR카드 PDF bytes)
    - issues: 이름 누락 등으로 제외된 셀에 대한 경고 목록 (페이지, 학번(있으면), 사유)
    """
    src = fitz.open(stream=pdf_bytes, filetype="pdf")
    students = []
    issues = []
    for page_no in range(len(src)):
        page = src[page_no]
        blocks = page.get_text("dict")["blocks"]
        if len(blocks) % 2 != 0:
            raise ValueError(f"페이지 {page_no}: 블록 개수가 홀수입니다 ({len(blocks)})")

        for i in range(len(blocks) // 2):
            text_block, img_block = blocks[2 * i], blocks[2 * i + 1]
            if text_block["type"] != 0 or img_block["type"] != 1:
                raise ValueError(f"페이지 {page_no}, 셀 {i}: 예상치 못한 블록 순서")

            student_id, name, reason = _parse_cell_text(text_block)
            if reason is not None:
                issues.append({"page": page_no + 1, "id": student_id, "reason": reason})
                continue

            tb, ib = text_block["bbox"], img_block["bbox"]
            clip = fitz.Rect(
                min(tb[0], ib[0]) - CROP_PAD,
                tb[1] - CROP_PAD,
                max(tb[2], ib[2]) + CROP_PAD,
                ib[3] + CROP_PAD,
            )

            crop_doc = fitz.open()
            crop_page = crop_doc.new_page(width=clip.width, height=clip.height)
            crop_page.show_pdf_page(crop_page.rect, src, page_no, clip=clip)
            crop_pdf_bytes = crop_doc.tobytes()
            crop_doc.close()

            students.append(
                {
                    "id": student_id,
                    "name": name,
                    "grade": int(student_id[0]),
                    "cls": int(student_id[1:3]),
                    "num": int(student_id[3:5]),
                    "crop_pdf_bytes": crop_pdf_bytes,
                }
            )
    src.close()
    return students, issues


def render_png_from_pdf_bytes(pdf_bytes, zoom=6):
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    pix = doc[0].get_pixmap(matrix=fitz.Matrix(zoom, zoom))
    return pix.tobytes("png")


def replace_placeholders(paragraph, values):
    full_text = paragraph.text
    if "{{" not in full_text:
        return
    new_text = PLACEHOLDER_RE.sub(lambda m: str(values[m.group(1)]), full_text)

    first_run = paragraph.runs[0] if paragraph.runs else None
    font_name = first_run.font.name if first_run else None
    size = first_run.font.size if first_run else None
    bold = first_run.font.bold if first_run else None
    italic = first_run.font.italic if first_run else None
    color = None
    if first_run is not None and first_run.font.color is not None and first_run.font.color.type is not None:
        color = first_run.font.color.rgb

    for r in list(paragraph.runs):
        r._element.getparent().remove(r._element)

    run = paragraph.add_run(new_text)
    if font_name:
        run.font.name = font_name
    if size:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def find_image_paragraph(doc):
    candidates = [p for p in doc.paragraphs if p._element.findall(".//" + qn("w:drawing"))]
    if len(candidates) != 1:
        raise ValueError(f"템플릿에서 이미지가 있는 문단을 정확히 1개 찾지 못했습니다 (발견: {len(candidates)}개).")
    return candidates[0]


def replace_qr_image(doc, png_bytes, width_cm=DEFAULT_QR_WIDTH_CM):
    paragraph = find_image_paragraph(doc)

    blip = paragraph._element.find(".//" + qn("a:blip"))
    old_rid = blip.get(qn("r:embed")) if blip is not None else None

    for r in list(paragraph.runs):
        r._element.getparent().remove(r._element)
    if old_rid:
        doc.part.drop_rel(old_rid)

    run = paragraph.add_run()
    run.add_picture(io.BytesIO(png_bytes), width=Cm(width_cm))


def build_student_docx(template_source, student, width_cm=DEFAULT_QR_WIDTH_CM):
    """template_source: 파일 경로 또는 docx 원본 bytes. student: parse_source_pdf가 반환한 dict 1개."""
    if isinstance(template_source, (bytes, bytearray)):
        template_source = io.BytesIO(template_source)
    doc = Document(template_source)

    values = {"이름": student["name"], "학번": student["id"], "반": student["cls"], "번호": student["num"]}
    for p in doc.paragraphs:
        replace_placeholders(p, values)

    png_bytes = render_png_from_pdf_bytes(student["crop_pdf_bytes"])
    replace_qr_image(doc, png_bytes, width_cm=width_cm)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def merge_docx_bytes_list(docx_bytes_list):
    if not docx_bytes_list:
        raise ValueError("병합할 문서가 없습니다.")

    master = Document(io.BytesIO(docx_bytes_list[0]))
    composer = Composer(master)

    for b in docx_bytes_list[1:]:
        master.add_page_break()
        composer.append(Document(io.BytesIO(b)))

    buf = io.BytesIO()
    composer.save(buf)
    return buf.getvalue()


def student_filename(student, ext):
    return f"{student['id']}_{student['name']}({student['grade']}학년).{ext}"


def find_template_placeholders(template_bytes):
    """템플릿 안에 실제로 쓰인 {{...}} 필드 목록을 중복 없이 반환 (미리보기/검증용)."""
    doc = Document(io.BytesIO(template_bytes))
    found = set()
    for p in doc.paragraphs:
        found.update(re.findall(r"\{\{(.+?)\}\}", p.text))
    return sorted(found)
